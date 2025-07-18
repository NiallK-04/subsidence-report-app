import streamlit as st
import requests
import folium
from streamlit_folium import st_folium
from docx import Document
from datetime import date
import io
import os
from PIL import Image
from selenium import webdriver
import chromedriver_autoinstaller
import time
import docx.shared

# Ensure Chromedriver is available
chromedriver_autoinstaller.install()

# --- Helper functions ---
def get_coords_from_eircode(eircode, api_key):
    url = f"https://api.opencagedata.com/geocode/v1/json?q={eircode}&key={api_key}&countrycode=ie&limit=1"
    response = requests.get(url)
    if response.status_code == 200:
        results = response.json()['results']
        if results:
            lat = results[0]['geometry']['lat']
            lng = results[0]['geometry']['lng']
            return lat, lng
    return None, None

def query_gsi_geology(lat, lon):
    url = "https://secure.dccae.gov.ie/gsi-wfs/public"
    params = {
        "service": "WFS",
        "version": "1.1.0",
        "request": "GetFeature",
        "typeName": "bedrock100k",
        "srsName": "EPSG:4326",
        "outputFormat": "application/json",
        "bbox": f"{lon},{lat},{lon},{lat}"
    }
    try:
        response = requests.get(url, params=params)
        data = response.json()
        if data['features']:
            rock = data['features'][0]['properties']['ROCKNAME']
            return f"The underlying bedrock geology is primarily {rock.lower()}."
    except:
        pass
    return "Geological information unavailable."

def capture_map_snapshot(lat, lon, map_html_path, image_path):
    fmap = folium.Map(location=[lat, lon], zoom_start=16)
    folium.Marker([lat, lon], tooltip="Property").add_to(fmap)
    fmap.save(map_html_path)

    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    driver = webdriver.Chrome(options=options)
    driver.set_window_size(800, 600)
    driver.get(f'file://{map_html_path}')
    time.sleep(2)
    driver.save_screenshot(image_path)
    driver.quit()

# --- Streamlit UI ---
st.set_page_config(page_title="Forensic Subsidence Report Assistant", layout="centered")
st.title("🛠️ Forensic Subsidence Report Assistant (2025)")

with st.form("subsidence_form"):
    insurer = st.text_input("Insurer Name")
    claim_ref = st.text_input("Claim Reference")
    address = st.text_area("Property Address")
    inspection_date = st.date_input("Date of Inspection", value=date.today())

    eircode = st.text_input("Eircode")

    historical_photos = st.file_uploader("Upload Historical Photos (optional)", accept_multiple_files=True, type=['png', 'jpg', 'jpeg'])

    submit = st.form_submit_button("Generate Report")

if submit:
    api_key = st.secrets.get("OPENCAGE_API_KEY", "")
    if not api_key:
        st.error("API key not found in Streamlit secrets. Please add OPENCAGE_API_KEY in your app settings.")
    else:
        lat, lon = get_coords_from_eircode(eircode.strip(), api_key)
        if not lat:
            st.error("Unable to resolve Eircode.")
        else:
            st.success(f"Location resolved: {lat:.5f}, {lon:.5f}")

            geology_summary = query_gsi_geology(lat, lon)

            map_html = "/tmp/temp_map.html"
            map_image = "/tmp/map_image.png"
            capture_map_snapshot(lat, lon, map_html, map_image)

            doc = Document()
            doc.add_heading("Subsidence Report", 0)

            doc.add_heading("1. Property & Claim Info", level=1)
            doc.add_paragraph(f"Insurer: {insurer}")
            doc.add_paragraph(f"Claim Ref: {claim_ref}")
            doc.add_paragraph(f"Address: {address}")
            doc.add_paragraph(f"Eircode: {eircode}")
            doc.add_paragraph(f"Coordinates: {lat:.5f}, {lon:.5f}")
            doc.add_paragraph(f"Inspection Date: {inspection_date.strftime('%d %B %Y')}")

            doc.add_heading("2. Site Location Map", level=1)
            doc.add_picture(map_image, width=docx.shared.Inches(5.5))

            doc.add_heading("3. Geological Summary", level=1)
            doc.add_paragraph(geology_summary)

            if historical_photos:
                doc.add_heading("4. Historical Photos", level=1)
                for idx, photo in enumerate(historical_photos, start=1):
                    image = Image.open(photo)
                    temp_path = f"/tmp/{photo.name}"
                    image.save(temp_path)
                    doc.add_paragraph(f"Figure {idx}: {photo.name}")
                    doc.add_picture(temp_path, width=docx.shared.Inches(5.5))

            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("📄 Download Word Report", data=buf.getvalue(), file_name="subsidence_report.docx")

            st.image(map_image, caption="Map Snapshot", use_column_width=True)
            st.info(geology_summary)